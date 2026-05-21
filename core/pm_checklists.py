from __future__ import annotations

import json
import time
from pathlib import Path
from typing import Any

from .analysis_common import timestamp_for_report
from .constants import TOOLKIT_ROOT
from .logging import log_tool_run
from .paths import resolve_project_paths
from .result import ToolResult
from .safe_files import ensure_directory, safe_write_text
from .workbook_io import row_dicts


TOOL_ID = "pm_checklist_generator"
TOOL_NAME = "EOAT PM Checklist Generator"

BASE_CHECKS = [
    "Inspect mounting hardware for looseness, missing hardware, wear, or damage.",
    "Verify EOAT alignment condition and note any need for verification.",
    "Inspect cable management for loose, rubbing, pinched, or damaged cables.",
    "Check quick disconnect fittings for damage, looseness, labeling, and repeatability.",
    "Look for sharp edges, pinch points, loose fasteners, and other safety concerns.",
    "Verify documentation/BOM/process binder status and note missing references.",
]

TYPE_CHECKS = {
    "vacuum": [
        "Inspect vacuum cups for wear, cracking, deformation, contamination, and loss of sealing ability.",
        "Verify correct vacuum cup type/material and cup diameter/size are installed.",
        "Verify vacuum cup mounting hardware is secure.",
        "Inspect vacuum tubing for kinks, leaks, abrasion, poor bend radius, or messy routing.",
        "Verify vacuum generator/ejector condition if documented.",
        "Verify vacuum zones are labeled and functioning if applicable.",
        "Verify vacuum confirmation sensor operation if present.",
        "Review part drop or mis-pick history and note recurring conditions.",
    ],
    "mechanical": [
        "Inspect gripper fingers/jaws for wear, damage, looseness, or contamination.",
        "Check actuator motion and verify open/close function.",
        "Inspect pivot points, slides, bushings, and other wear areas.",
        "Check gripper mounting hardware and locking hardware.",
        "Verify part-present sensor operation if present.",
        "Check contact pads or part-contact surfaces.",
    ],
    "hybrid": [
        "Complete all vacuum checks.",
        "Complete all mechanical gripper checks.",
        "Verify coordination between vacuum and mechanical gripping systems.",
        "Verify sensor logic/confirmation method if documented.",
    ],
}

GENERIC_TYPES = ["Vacuum", "Mechanical Gripper", "Hybrid"]


def _clean(value: Any) -> str:
    return str(value or "").strip()


def _kind(eoat_type: str) -> str:
    text = eoat_type.lower()
    if "hybrid" in text:
        return "hybrid"
    if "mechanical" in text or "gripper" in text:
        return "mechanical"
    if "vacuum" in text:
        return "vacuum"
    return "vacuum"


def _slug(text: str, fallback: str = "All_EOATs") -> str:
    cleaned = "".join(char if char.isalnum() else "_" for char in text.strip())
    while "__" in cleaned:
        cleaned = cleaned.replace("__", "_")
    return cleaned.strip("_") or fallback


def _load_template_checks(eoat_type: str) -> list[str]:
    template_path = TOOLKIT_ROOT / "data_templates" / "pm_checklist_templates.json"
    if not template_path.exists():
        return []
    try:
        data = json.loads(template_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return []
    for template in data.get("templates", []):
        if _clean(template.get("eoat_type")).lower() == eoat_type.lower():
            return [str(check) for check in template.get("checks", [])]
    return []


def _checks_for_type(eoat_type: str) -> list[str]:
    kind = _kind(eoat_type)
    checks = list(BASE_CHECKS)
    if kind == "hybrid":
        checks.extend(TYPE_CHECKS["vacuum"])
        checks.extend(TYPE_CHECKS["mechanical"])
        checks.extend(TYPE_CHECKS["hybrid"])
    else:
        checks.extend(TYPE_CHECKS[kind])
    checks.extend(_load_template_checks(eoat_type))
    deduped: list[str] = []
    seen: set[str] = set()
    for check in checks:
        key = check.lower()
        if key not in seen:
            deduped.append(check)
            seen.add(key)
    return deduped


def _missing_data_warnings(row: dict[str, Any]) -> list[str]:
    fields = [
        "EOAT Type",
        "EOAT Moves",
        "Robot Type",
        "Tubing Condition",
        "Cable Management Condition",
        "Mounting Hardware Condition",
        "EOAT Alignment Condition",
        "Known Issues",
        "Maintenance Frequency",
        "Drawing/CAD Available?",
        "BOM Available?",
        "Process Binder Complete?",
    ]
    missing = [field for field in fields if not _clean(row.get(field))]
    if not missing:
        return []
    press = _clean(row.get("Press/Machine #") or row.get("Audit ID") or "selected EOAT")
    return [f"{press}: missing data for {', '.join(missing)}."]


def build_pm_checklist_markdown(row: dict[str, Any] | None = None, eoat_type: str = "Vacuum") -> tuple[str, list[str]]:
    row = row or {}
    type_name = _clean(row.get("EOAT Type")) or eoat_type
    press = _clean(row.get("Press/Machine #")) or "Generic"
    audit_id = _clean(row.get("Audit ID")) or "N/A"
    title = f"PM Checklist - {press} ({type_name})" if press != "Generic" else f"Generic PM Checklist - {type_name}"
    checks = _checks_for_type(type_name)
    warnings = _missing_data_warnings(row) if row else []

    lines = [
        f"# {title}",
        "",
        f"- Audit ID: {audit_id}",
        f"- Plant/Area: {_clean(row.get('Plant/Area')) or 'N/A'}",
        f"- Robot Type: {_clean(row.get('Robot Type')) or 'N/A'}",
        f"- EOAT Type: {type_name}",
        f"- EOAT Moves: {_clean(row.get('EOAT Moves')) or 'N/A'}",
        f"- Known Issues: {_clean(row.get('Known Issues')) or 'None documented'}",
        "",
        "## Checklist",
        "| Item | Pass/Fail | Comments | Recommended Follow-Up |",
        "| --- | --- | --- | --- |",
    ]
    for check in checks:
        lines.append(f"| {check} |  |  |  |")
    lines.extend(
        [
            "",
            "## Documentation Check",
            f"- Spare Parts Identified?: {_clean(row.get('Spare Parts Identified?')) or 'Unknown'}",
            f"- Drawing/CAD Available?: {_clean(row.get('Drawing/CAD Available?')) or 'Unknown'}",
            f"- BOM Available?: {_clean(row.get('BOM Available?')) or 'Unknown'}",
            f"- Process Binder Complete?: {_clean(row.get('Process Binder Complete?')) or 'Unknown'}",
            "",
            "## Notes",
            _clean(row.get("Notes")) or "Add notes during PM inspection.",
        ]
    )
    if warnings:
        lines.extend(["", "## Missing Data Warnings"])
        lines.extend(f"- {warning}" for warning in warnings)
    return "\n".join(lines) + "\n", warnings


def _filter_rows(rows: list[dict[str, Any]], audit_id: str | None, press: str | None, all_audited: bool) -> list[dict[str, Any]]:
    if audit_id:
        return [row for row in rows if _clean(row.get("Audit ID")).lower() == audit_id.strip().lower()]
    if press:
        return [row for row in rows if _clean(row.get("Press/Machine #")).lower() == press.strip().lower()]
    if all_audited:
        return [row for row in rows if _clean(row.get("Status")).lower() in {"audited", "needs follow-up", "candidate for pilot"}]
    return []


def _write_docx_if_requested(path: Path, markdown: str) -> Path | None:
    try:
        from docx import Document
    except Exception:
        return None
    doc = Document()
    for line in markdown.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|"):
            doc.add_paragraph(line)
        elif line.strip():
            doc.add_paragraph(line)
    docx_path = path.with_suffix(".docx")
    doc.save(docx_path)
    return docx_path


def _unique_markdown_path(folder: Path, stem: str, stamp: str) -> Path:
    path = folder / f"{stem}_{stamp}.md"
    if not path.exists():
        return path
    for index in range(2, 1000):
        candidate = folder / f"{stem}_{stamp}_{index}.md"
        if not candidate.exists():
            return candidate
    return folder / f"{stem}_{int(time.time())}.md"


def generate_pm_checklists(
    project_root: str | Path,
    audit_id: str | None = None,
    press: str | None = None,
    all_audited: bool = False,
    generic: bool = False,
    formats: list[str] | None = None,
) -> ToolResult:
    start = time.perf_counter()
    paths = resolve_project_paths(project_root)
    output_dir = ensure_directory(paths.pm_generated_checklists)
    formats = [fmt.lower() for fmt in (formats or ["markdown"])]
    files_created: list[str] = []
    details: list[str] = []
    warnings: list[str] = []
    rows: list[dict[str, Any]] = []

    if paths.master_workbook.exists():
        try:
            rows = row_dicts(paths.master_workbook, "EOAT Inventory")
        except Exception as exc:
            warnings.append(f"Could not read EOAT Inventory; generic templates are still available: {exc}")
    else:
        warnings.append(f"Master workbook not found: {paths.master_workbook}")

    selected_rows = _filter_rows(rows, audit_id, press, all_audited)
    if not selected_rows:
        generic = True
        if audit_id or press or all_audited:
            warnings.append("No matching EOAT rows found, so generic templates were generated.")

    stamp = timestamp_for_report()
    if generic:
        for type_name in GENERIC_TYPES:
            markdown, row_warnings = build_pm_checklist_markdown(None, eoat_type=type_name)
            warnings.extend(row_warnings)
            path = safe_write_text(_unique_markdown_path(output_dir, f"PM_Checklist_Generic_{_slug(type_name)}", stamp), markdown, overwrite=False)
            files_created.append(str(path))
            if "docx" in formats:
                docx = _write_docx_if_requested(path, markdown)
                if docx:
                    files_created.append(str(docx))
                else:
                    warnings.append("DOCX output requested, but python-docx is unavailable.")
        details.append("Generated generic PM checklist templates.")
    else:
        for row in selected_rows:
            markdown, row_warnings = build_pm_checklist_markdown(row)
            warnings.extend(row_warnings)
            name = _slug(_clean(row.get("Press/Machine #")) or _clean(row.get("Audit ID")) or "EOAT")
            path = safe_write_text(_unique_markdown_path(output_dir, f"PM_Checklist_{name}", stamp), markdown, overwrite=False)
            files_created.append(str(path))
            if "docx" in formats:
                docx = _write_docx_if_requested(path, markdown)
                if docx:
                    files_created.append(str(docx))
                else:
                    warnings.append("DOCX output requested, but python-docx is unavailable.")
        details.append(f"Generated {len(selected_rows)} EOAT-specific PM checklist(s).")

    if "excel" in formats:
        warnings.append("Excel checklist output is planned for a later refinement; Markdown was generated.")

    result = ToolResult.ok(
        TOOL_ID,
        TOOL_NAME,
        f"Generated {len(files_created)} checklist file(s).",
        details=details,
        warnings=warnings,
        files_created=files_created,
        output_reports=files_created,
        metrics={"checklist_files": len(files_created), "inventory_rows_read": len(rows)},
        duration_seconds=time.perf_counter() - start,
    )
    warning = log_tool_run(result, project_root)
    if warning:
        result.warnings.append(warning)
    return result
