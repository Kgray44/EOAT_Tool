from __future__ import annotations

import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from .analysis_common import count_by, parse_score, table_from_rows
from .paths import resolve_project_paths
from .reports import list_recent_files
from .safe_files import ensure_directory
from .workbook_io import row_dicts


def stamp() -> str:
    return time.strftime("%Y-%m-%d_%H%M")


def clean(value: Any) -> str:
    return str(value or "").strip()


def unique_package_dir(parent: str | Path, base_name: str) -> Path:
    root = ensure_directory(parent)
    candidate = root / f"{base_name}_{stamp()}"
    if not candidate.exists():
        return ensure_directory(candidate)
    for index in range(2, 1000):
        indexed = root / f"{base_name}_{stamp()}_{index}"
        if not indexed.exists():
            return ensure_directory(indexed)
    return ensure_directory(root / f"{base_name}_{int(time.time())}")


def safe_rows(project_root: str | Path, sheet_name: str) -> tuple[list[dict[str, Any]], str | None]:
    paths = resolve_project_paths(project_root)
    if not paths.master_workbook.exists():
        return [], f"Master workbook not found: {paths.master_workbook}"
    try:
        return row_dicts(paths.master_workbook, sheet_name), None
    except Exception as exc:
        return [], f"Could not read {sheet_name}: {exc}"


def workbook_metrics(project_root: str | Path) -> tuple[dict[str, Any], list[str]]:
    warnings: list[str] = []
    inventory, warning = safe_rows(project_root, "EOAT Inventory")
    if warning:
        warnings.append(warning)
    issues, warning = safe_rows(project_root, "Issue Log")
    if warning:
        warnings.append(warning)
    interviews, warning = safe_rows(project_root, "Interview Notes")
    if warning:
        warnings.append(warning)
    photos, warning = safe_rows(project_root, "Photo Index")
    if warning:
        warnings.append(warning)
    pilots, warning = safe_rows(project_root, "Pilot Candidates")
    if warning:
        warnings.append(warning)
    fmea, warning = safe_rows(project_root, "FMEA Draft")
    if warning:
        warnings.append(warning)
    kpis, warning = safe_rows(project_root, "KPI Baseline")
    if warning:
        warnings.append(warning)

    audited = sum(1 for row in inventory if clean(row.get("Status")).lower() in {"audited", "needs follow-up", "candidate for pilot"})
    pilot_flagged = sum(1 for row in inventory if clean(row.get("Pilot Candidate?")).lower() in {"yes", "maybe"})
    metrics = {
        "Total EOATs identified": len(inventory),
        "Total EOATs audited": audited,
        "Issues logged": len(issues),
        "Interviews logged": len(interviews),
        "Photos indexed": len(photos),
        "Pilot candidates flagged": pilot_flagged + len(pilots),
        "FMEA risks identified": len(fmea),
        "KPI records available": len(kpis),
    }
    return metrics, warnings


def top_issue_categories(project_root: str | Path, limit: int = 5) -> list[dict[str, Any]]:
    issues, _warning = safe_rows(project_root, "Issue Log")
    counts = count_by(issues, "Issue Category")
    return [{"Issue Category": key, "Count": value} for key, value in sorted(counts.items(), key=lambda item: (-item[1], item[0]))[:limit]]


def top_fmea_risks(project_root: str | Path, limit: int = 5) -> list[dict[str, Any]]:
    rows, _warning = safe_rows(project_root, "FMEA Draft")
    ranked: list[dict[str, Any]] = []
    for row in rows:
        rpn = parse_score(row.get("RPN"))
        if rpn is None:
            sev = parse_score(row.get("Severity")) or 0
            freq = parse_score(row.get("Frequency")) or 0
            det = parse_score(row.get("Detectability")) or 0
            rpn = sev * freq * det if sev and freq and det else 0
        ranked.append(
            {
                "Press/Machine #": clean(row.get("Press/Machine #")),
                "Failure Mode": clean(row.get("Failure Mode")),
                "RPN": rpn,
                "Recommended Action": clean(row.get("Recommended Action")),
            }
        )
    return sorted(ranked, key=lambda row: int(row.get("RPN") or 0), reverse=True)[:limit]


def recent_report_map(project_root: str | Path) -> dict[str, list[Path]]:
    paths = resolve_project_paths(project_root)
    return {
        "Audit Progress": list_recent_files(paths.audit_progress_reports, 5),
        "Issue Analysis": list_recent_files(paths.issue_analysis_reports, 5),
        "Documentation Gaps": list_recent_files(paths.documentation_gap_reports, 5),
        "FMEA": list_recent_files(paths.fmea_reports, 5),
        "Pilot Candidates": list_recent_files(paths.pilot_project / "Candidate_Cells", 5),
        "KPI": list_recent_files(paths.kpi_dashboard_exports, 5),
        "PM Checklists": list_recent_files(paths.pm_generated_checklists, 5),
        "BOM": list_recent_files(paths.bom_standardization_reports, 5),
        "Weekly": list_recent_files(paths.weekly_reports, 5),
        "Mentor Briefs": list_recent_files(paths.mentor_briefs, 5),
    }


@dataclass(frozen=True)
class DeliverableStatus:
    name: str
    status: str
    evidence: list[str]
    notes: str


def status_table_markdown(statuses: list[DeliverableStatus]) -> list[str]:
    return table_from_rows(
        [{"Deliverable": item.name, "Status": item.status, "Evidence": "; ".join(item.evidence[:3]), "Notes": item.notes} for item in statuses],
        ["Deliverable", "Status", "Evidence", "Notes"],
    )


def metrics_markdown(metrics: dict[str, Any]) -> list[str]:
    return [f"- {key}: {value}" for key, value in metrics.items()]


def report_references_markdown(report_map: dict[str, list[Path]]) -> list[str]:
    lines: list[str] = []
    for label, files in report_map.items():
        if files:
            lines.append(f"- {label}: {files[0].name}")
        else:
            lines.append(f"- {label}: Not available yet")
    return lines


def make_simple_docx(markdown_path: Path, markdown_text: str) -> Path | None:
    try:
        from docx import Document
    except Exception:
        return None
    doc = Document()
    for line in markdown_text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|") or line.strip():
            doc.add_paragraph(line)
    output = markdown_path.with_suffix(".docx")
    doc.save(output)
    return output
