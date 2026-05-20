"""
Generate one EOAT daily status summary at a time.

This tool combines:
- Manual notes from the intern
- Local Git activity when the project is a Git repository
- Local file snapshot comparisons when Git is not available
- Week 1 schedule/task progress tracking

No internet access or cloud APIs are used.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
from datetime import date, datetime, time
from pathlib import Path
from typing import Any, Iterable


SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_PROJECT_ROOT = SCRIPT_DIR / "EOAT_Standardization_Project"

EXPECTED_WORKBOOK_SHEETS = [
    "EOAT Inventory",
    "Issue Log",
    "KPI Baseline",
    "Interview Notes",
    "Pilot Candidates",
    "FMEA Draft",
    "Action Items",
    "Photo Index",
]

IMPORTANT_INVENTORY_COLUMNS = [
    "Audit ID",
    "Audit Date",
    "Auditor",
    "Plant/Area",
    "Press/Machine #",
    "Robot Type",
    "EOAT Type",
    "Sensors Present?",
    "Tubing Condition",
    "Cable Management Condition",
    "Known Issues",
    "Maintenance Frequency",
    "Photos Taken?",
    "Status",
    "Priority",
    "Pilot Candidate?",
    "Notes",
]

WEEK1_DEFAULT_DAYS = {
    "1": [
        "Review EOAT project scope and deliverables",
        "Create project folder/document structure",
        "Start EOAT audit database",
        "Identify initial stakeholder questions",
        "Begin list of target robot cells",
    ],
    "2": [
        "Finalize audit template",
        "Begin target cell list",
        "Start first walkthrough/audit if approved",
        "Decide photo naming system",
        "Confirm initial audit priorities with mentor or supervisor",
    ],
    "3": [
        "Audit 2-4 accessible EOATs in detail",
        "Capture overall EOAT photos",
        "Capture vacuum cup/gripper photos",
        "Capture tubing routing photos",
        "Capture sensor and quick disconnect photos",
        "Interview operator or technician for each audited cell",
        "Clean up notes and link photos in the workbook",
    ],
    "4": [
        "Improve the audit template based on real floor use",
        "Audit another 3-5 EOATs",
        "Include at least one good example cell if possible",
        "Start Issue Log tab",
        "Group early issues into categories",
        "Flag possible pilot candidate cells",
    ],
    "5": [
        "Clean and normalize the audit database",
        "Organize photo folders",
        "Create Week 1 summary",
        "Review progress with mentor or supervisor",
        "Confirm Week 2 priorities",
        "Identify data gaps and blockers",
    ],
}

DAY1_COMPLETED = [
    "Reviewed EOAT project scope and deliverables",
    "Created project folder/document structure",
    "Started EOAT audit database",
    "Identified initial stakeholder questions",
    "Began list of target robot cells",
]

DAY1_NEED = [
    "Confirm priority robot cells",
    "Locate downtime/scrap/cycle time data sources",
    "Confirm photo/documentation rules",
    "Identify technicians/operators to shadow",
]

IGNORE_DIRS = {".git", "__pycache__", ".venv", "venv", "node_modules", ".activity_snapshots"}
IGNORE_FILES = {".DS_Store", "Thumbs.db"}
TEXT_SNIPPET_EXTS = {".md", ".txt", ".py", ".ps1", ".csv", ".json", ".yaml", ".yml"}
SECRET_NAME_PARTS = {"secret", "token", "key", "password", "credential", "auth"}
HASH_SIZE_LIMIT = 2 * 1024 * 1024
SNIPPET_SIZE_LIMIT = 200 * 1024
ALLOWED_STATUSES = ["Not started", "In progress", "Blocked", "Complete", "Skipped"]


def parse_args() -> argparse.Namespace:
    """Read command-line arguments."""
    parser = argparse.ArgumentParser(description="Create one EOAT daily status report.")
    parser.add_argument("--project-root", default=str(DEFAULT_PROJECT_ROOT), help="Path to EOAT_Standardization_Project.")
    parser.add_argument("--week", type=int, default=1, help="Project week number. Defaults to 1.")
    parser.add_argument("--day", type=int, help="Project day number, such as 1 or 2.")
    parser.add_argument("--date", dest="report_date", default=date.today().isoformat(), help="Report date in YYYY-MM-DD format.")
    parser.add_argument("--interactive", action="store_true", help="Force interactive prompts.")
    parser.add_argument("--include-git", action="store_true", help="Include Git status/log/diff summary if available.")
    parser.add_argument("--include-snapshot", action="store_true", help="Include local file snapshot comparison.")
    parser.add_argument("--init-snapshot", action="store_true", help="Create a baseline snapshot without generating a report.")
    parser.add_argument("--output-format", choices=["markdown", "docx"], default="markdown", help="Report format. Defaults to markdown.")
    parser.add_argument("--include-diff-snippets", action="store_true", help="Include tiny snippets from changed safe text files.")
    parser.add_argument("--completed", nargs="*", help="Manual completed items.")
    parser.add_argument("--need", nargs="*", help="Manual needed items.")
    parser.add_argument("--plan", nargs="*", help="Manual next-day plan items to append.")
    parser.add_argument("--note", nargs="*", help="Optional manual notes.")
    parser.add_argument("--use-day1-defaults", action="store_true", help="Use suggested Day 1 defaults.")
    parser.add_argument("--yes", action="store_true", help="Allow overwrite without prompting. Useful for automation.")
    parser.add_argument("--docx", action="store_true", help="Compatibility alias for --output-format docx.")
    return parser.parse_args()


def ask(prompt: str, default: str | None = None) -> str:
    """Prompt with an optional default value."""
    suffix = f" [{default}]" if default else ""
    answer = input(f"{prompt}{suffix}: ").strip()
    return answer or (default or "")


def ask_yes_no(prompt: str, default: bool = True) -> bool:
    """Ask a yes/no question."""
    default_text = "Y/n" if default else "y/N"
    answer = input(f"{prompt} ({default_text}): ").strip().lower()
    if not answer:
        return default
    return answer.startswith("y")


def ask_int(prompt: str, default: int | None = None) -> int:
    """Prompt until the user enters a valid integer."""
    while True:
        raw = ask(prompt, str(default) if default is not None else None)
        try:
            return int(raw)
        except ValueError:
            print("Please enter a whole number, such as 1 or 2.")


def validate_date(value: str) -> str:
    """Validate YYYY-MM-DD dates."""
    datetime.strptime(value, "%Y-%m-%d")
    return value


def print_items(title: str, items: Iterable[str]) -> None:
    """Display a short bullet list."""
    print(f"\n{title}:")
    for item in items:
        print(f"- {item}")


def normalize_list(values: list[str] | None) -> list[str]:
    """Clean a command-line list."""
    if not values:
        return []
    return [value.strip() for value in values if value.strip()]


def dedupe(items: Iterable[str]) -> list[str]:
    """Keep list order while removing duplicates."""
    seen = set()
    result = []
    for item in items:
        key = item.strip().lower()
        if item.strip() and key not in seen:
            seen.add(key)
            result.append(item.strip())
    return result


def slug(text: str) -> str:
    """Normalize text for light task matching."""
    return "".join(char.lower() if char.isalnum() else " " for char in text)


def keyword_match(task_text: str, evidence_text: str) -> bool:
    """Simple local-only matching between tasks and evidence/manual notes."""
    task_words = {word[:5] for word in slug(task_text).split() if len(word) > 3}
    evidence_words = {word[:5] for word in slug(evidence_text).split() if len(word) > 3}
    if not task_words or not evidence_words:
        return False
    return len(task_words & evidence_words) >= min(2, len(task_words))


def admin_dir(project_root: Path) -> Path:
    """Return project admin directory."""
    return project_root / "00_Project_Admin"


def report_dir(project_root: Path) -> Path:
    """Return daily reports directory."""
    return admin_dir(project_root) / "Daily_Status_Reports"


def snapshot_dir(project_root: Path) -> Path:
    """Return snapshot storage directory."""
    return report_dir(project_root) / ".activity_snapshots"


def schedule_path(project_root: Path, week: int) -> Path:
    """Return schedule JSON path for a week."""
    return admin_dir(project_root) / f"project_schedule_week{week}.json"


def progress_path(project_root: Path, week: int) -> Path:
    """Return task progress JSON path for a week."""
    return admin_dir(project_root) / f"task_progress_week{week}.json"


def load_json(path: Path, default: dict) -> dict:
    """Load JSON with a safe default."""
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        backup = path.with_name(f"{path.stem}_invalid_{datetime.now().strftime('%Y%m%d_%H%M%S')}{path.suffix}")
        path.rename(backup)
        return default


def save_json(path: Path, data: dict) -> None:
    """Write pretty JSON."""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, indent=2), encoding="utf-8")


def schedule_file_matches_week(path: Path, week: int) -> bool:
    """Return True when a help document appears to belong to the requested week."""
    name = path.stem.lower()
    compact = re.sub(r"[^a-z0-9]+", "", name)
    return f"week{week}" in compact and "schedule" in compact


def schedule_help_files_for_week(project_root: Path, week: int) -> list[Path]:
    """Find local schedule help documents for a week."""
    help_dir = project_root.parent / "Project_Help_Documents"
    if not help_dir.exists():
        return []
    return [
        path
        for path in help_dir.iterdir()
        if path.is_file() and schedule_file_matches_week(path, week)
    ]


def extract_schedule_from_pdf(path: Path) -> dict[str, list[str]]:
    """Extract Dn-Tn planned task lines from a weekly schedule PDF when possible."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return {}

    try:
        text = "\n".join(page.extract_text() or "" for page in PdfReader(str(path)).pages)
    except Exception:
        return {}

    days: dict[str, list[str]] = {}
    pattern = re.compile(r"D(\d+)-T\d+\s*[-\u2013\u2014]\s*(.+)")
    for raw_line in text.splitlines():
        line = raw_line.strip()
        match = pattern.match(line)
        if not match:
            continue
        day, task_text = match.groups()
        task_text = task_text.strip().replace("\u2013", "-").replace("\u2014", "-")
        days.setdefault(day, [])
        if task_text not in days[day]:
            days[day].append(task_text)
    return days


def default_schedule_for_week(week: int) -> dict:
    """Return a generic schedule object for any project week."""
    if week == 1:
        days = WEEK1_DEFAULT_DAYS
        source = "Built-in Week 1 fallback schedule"
    else:
        days = {str(day): [] for day in range(1, 6)}
        source = "Template schedule; add a matching weekly schedule file in Project_Help_Documents"
    return {"week": week, "source": source, "source_files": [], "days": days}


def build_schedule_from_help_files(project_root: Path, week: int) -> dict:
    """Build a schedule from Project_Help_Documents when a matching file exists."""
    schedule = default_schedule_for_week(week)
    help_files = schedule_help_files_for_week(project_root, week)
    schedule["source_files"] = [str(path.relative_to(project_root.parent)) for path in help_files]
    for help_file in help_files:
        if help_file.suffix.lower() == ".json":
            loaded = load_json(help_file, {})
            if loaded.get("days"):
                loaded["week"] = week
                loaded.setdefault("source_files", schedule["source_files"])
                return loaded
        if help_file.suffix.lower() == ".pdf":
            extracted_days = extract_schedule_from_pdf(help_file)
            if extracted_days:
                schedule["days"] = extracted_days
                schedule["source"] = f"Extracted from {help_file.name}"
                break
    return schedule


def ensure_schedule(project_root: Path, week: int) -> dict:
    """Create or load the local schedule definition for the requested week."""
    path = schedule_path(project_root, week)
    default = build_schedule_from_help_files(project_root, week)
    if path.exists():
        schedule = load_json(path, default)
        existing_days = schedule.get("days", {})
        has_existing_tasks = any(existing_days.get(day) for day in existing_days)
        default_days = default.get("days", {})
        has_default_tasks = any(default_days.get(day) for day in default_days)
        if not existing_days or (not has_existing_tasks and has_default_tasks):
            schedule = default
            save_json(path, schedule)
        return schedule
    save_json(path, default)
    return default


def initial_progress(schedule: dict) -> dict:
    """Build initial task progress rows from the schedule."""
    tasks = []
    week = int(schedule.get("week", 1))
    for day_text, day_tasks in schedule.get("days", {}).items():
        day = int(day_text)
        for index, task_text in enumerate(day_tasks, start=1):
            tasks.append(
                {
                    "week": week,
                    "day": day,
                    "task_id": f"W{week}D{day}T{index}",
                    "task_text": task_text,
                    "status": "Not started",
                    "completed_date": "",
                    "evidence": [],
                    "notes": "",
                }
            )
    return {"tasks": tasks}


def ensure_progress(project_root: Path, week: int, schedule: dict) -> dict:
    """Create or load task progress."""
    path = progress_path(project_root, week)
    progress = load_json(path, initial_progress(schedule))
    existing_ids = {task["task_id"] for task in progress.get("tasks", [])}
    for task in initial_progress(schedule)["tasks"]:
        if task["task_id"] not in existing_ids:
            progress.setdefault("tasks", []).append(task)
    save_json(path, progress)
    return progress


def run_command(args: list[str], cwd: Path) -> tuple[int, str, str]:
    """Run a local command and capture output."""
    try:
        result = subprocess.run(args, cwd=cwd, capture_output=True, text=True, timeout=30, check=False)
        return result.returncode, result.stdout.strip(), result.stderr.strip()
    except (FileNotFoundError, subprocess.SubprocessError) as exc:
        return 1, "", str(exc)


def git_executable() -> str | None:
    """Find Git on PATH or in standard Git for Windows install locations."""
    found = shutil.which("git")
    if found:
        return found
    candidates = [
        Path.home() / "AppData/Local/Programs/Git/cmd/git.exe",
        Path("C:/Program Files/Git/cmd/git.exe"),
        Path("C:/Program Files (x86)/Git/cmd/git.exe"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def git_root(project_root: Path) -> Path | None:
    """Return the Git root if project_root is inside a Git repository."""
    git = git_executable()
    if not git:
        return None
    code, stdout, _ = run_command([git, "rev-parse", "--show-toplevel"], project_root)
    if code != 0 or not stdout:
        return None
    return Path(stdout)


def parse_git_status(short_status: str) -> dict[str, list[dict[str, str]]]:
    """Parse porcelain-ish short Git status into useful groups."""
    groups = {
        "created_files": [],
        "modified_files": [],
        "deleted_files": [],
        "renamed_files": [],
        "staged_files": [],
        "uncommitted_files": [],
    }
    for line in short_status.splitlines():
        if not line.strip():
            continue
        status = line[:2]
        path_text = line[3:].strip()
        status_text = "modified"
        target_path = path_text
        if " -> " in path_text:
            target_path = path_text.split(" -> ", 1)[1]
            status_text = "renamed"
        elif "A" in status or status == "??":
            status_text = "added"
        elif "D" in status:
            status_text = "deleted"
        elif "M" in status:
            status_text = "modified"

        item = {"path": target_path, "status": status_text, "git_status": status}
        if status_text == "added":
            groups["created_files"].append(item)
        elif status_text == "deleted":
            groups["deleted_files"].append(item)
        elif status_text == "renamed":
            groups["renamed_files"].append(item)
        else:
            groups["modified_files"].append(item)

        if status[0] not in (" ", "?"):
            groups["staged_files"].append(item)
        if status[1] != " " or status == "??":
            groups["uncommitted_files"].append(item)
    return groups


def parse_name_status(output: str) -> list[dict[str, str]]:
    """Parse Git name-status output."""
    items = []
    for line in output.splitlines():
        parts = line.split("\t")
        if len(parts) < 2:
            continue
        code = parts[0]
        path = parts[-1]
        status = {"A": "added", "M": "modified", "D": "deleted", "R": "renamed"}.get(code[:1], code)
        items.append({"path": path, "status": status, "git_status": code})
    return items


def collect_git_activity(project_root: Path, report_date: str) -> dict[str, Any]:
    """Collect concise Git activity without raw diffs."""
    git = git_executable()
    if not git:
        return {"available": False, "used": False, "note": "Git is not available; used file snapshot comparison if enabled."}
    root = git_root(project_root)
    if not root:
        return {"available": False, "used": False, "note": "Git repository not initialized; used file snapshot comparison if enabled."}

    day_start = datetime.combine(datetime.strptime(report_date, "%Y-%m-%d").date(), time.min).isoformat()
    day_end = datetime.combine(datetime.strptime(report_date, "%Y-%m-%d").date(), time.max).isoformat()
    commands = {
        "status_short": [git, "status", "--short"],
        "diff_stat": [git, "diff", "--stat"],
        "diff_name_status": [git, "diff", "--name-status"],
        "diff_numstat": [git, "diff", "--numstat"],
        "cached_stat": [git, "diff", "--cached", "--stat"],
        "cached_name_status": [git, "diff", "--cached", "--name-status"],
        "log_today": [git, "log", f"--since={day_start}", f"--until={day_end}", "--oneline", "--stat"],
        "log_today_oneline": [git, "log", f"--since={day_start}", f"--until={day_end}", "--oneline"],
    }
    outputs: dict[str, str] = {}
    for key, command in commands.items():
        code, stdout, _ = run_command(command, root)
        outputs[key] = stdout if code == 0 else ""

    parsed_status = parse_git_status(outputs["status_short"])
    unstaged_changes = parse_name_status(outputs["diff_name_status"])
    staged_changes = parse_name_status(outputs["cached_name_status"])
    commits = [line for line in outputs["log_today_oneline"].splitlines() if line.strip()]

    return {
        "available": True,
        "used": True,
        "git_root": str(root),
        "raw": outputs,
        "parsed_status": parsed_status,
        "unstaged_changes": unstaged_changes,
        "staged_changes": staged_changes,
        "commits_today": commits,
    }


def is_ignored(path: Path, relative: Path) -> bool:
    """Return True when a file should be ignored by snapshots."""
    if path.name in IGNORE_FILES or path.name.startswith("~$"):
        return True
    return any(part in IGNORE_DIRS for part in relative.parts)


def appears_binary(path: Path) -> bool:
    """Lightweight binary detection."""
    try:
        sample = path.read_bytes()[:2048]
    except OSError:
        return True
    if b"\x00" in sample:
        return True
    try:
        sample.decode("utf-8")
        return False
    except UnicodeDecodeError:
        return True


def sha256_if_small(path: Path, size: int) -> str:
    """Hash reasonably small files only."""
    if size > HASH_SIZE_LIMIT:
        return ""
    digest = hashlib.sha256()
    try:
        with path.open("rb") as file:
            for chunk in iter(lambda: file.read(65536), b""):
                digest.update(chunk)
    except OSError:
        return ""
    return digest.hexdigest()


def build_snapshot(project_root: Path) -> dict[str, Any]:
    """Create a local file activity snapshot."""
    files: dict[str, dict[str, Any]] = {}
    for root, dirnames, filenames in os.walk(project_root):
        root_path = Path(root)
        dirnames[:] = [name for name in dirnames if name not in IGNORE_DIRS]
        for filename in filenames:
            path = root_path / filename
            relative = path.relative_to(project_root)
            if is_ignored(path, relative):
                continue
            try:
                stat = path.stat()
            except OSError:
                continue
            size = stat.st_size
            binary = appears_binary(path)
            files[relative.as_posix()] = {
                "path": relative.as_posix(),
                "size": size,
                "modified_timestamp": stat.st_mtime,
                "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
                "sha256": sha256_if_small(path, size),
                "extension": path.suffix.lower(),
                "is_binary": binary,
                "is_large": size > HASH_SIZE_LIMIT,
            }
    return {"created_at": datetime.now().isoformat(timespec="seconds"), "root": str(project_root), "files": files}


def latest_snapshot_path(project_root: Path) -> Path:
    """Return latest snapshot path."""
    return snapshot_dir(project_root) / "latest_snapshot.json"


def save_snapshot(project_root: Path, snapshot: dict[str, Any], report_date: str) -> Path:
    """Save dated and latest snapshots."""
    directory = snapshot_dir(project_root)
    directory.mkdir(parents=True, exist_ok=True)
    dated = directory / f"snapshot_{report_date}_{datetime.now().strftime('%H%M%S')}.json"
    save_json(dated, snapshot)
    save_json(latest_snapshot_path(project_root), snapshot)
    return dated


def compare_snapshots(previous: dict[str, Any] | None, current: dict[str, Any]) -> dict[str, list[dict[str, Any]]]:
    """Compare two snapshots."""
    if not previous:
        created = [{"path": path, "status": "added", **meta} for path, meta in current.get("files", {}).items()]
        return {"created": created, "modified": [], "deleted": [], "renamed": [], "large_changed": []}

    old_files = previous.get("files", {})
    new_files = current.get("files", {})
    created_paths = set(new_files) - set(old_files)
    deleted_paths = set(old_files) - set(new_files)
    common_paths = set(new_files) & set(old_files)

    created = [{"path": path, "status": "added", **new_files[path]} for path in sorted(created_paths)]
    deleted = [{"path": path, "status": "deleted", **old_files[path]} for path in sorted(deleted_paths)]
    modified = []
    for path in sorted(common_paths):
        old = old_files[path]
        new = new_files[path]
        if old.get("sha256") and new.get("sha256"):
            changed = old["sha256"] != new["sha256"]
        else:
            changed = old.get("size") != new.get("size") or old.get("modified_timestamp") != new.get("modified_timestamp")
        if changed:
            modified.append({"path": path, "status": "modified", **new})

    renamed = []
    created_by_hash = {item.get("sha256"): item for item in created if item.get("sha256")}
    for deleted_item in deleted[:]:
        old_hash = deleted_item.get("sha256")
        if old_hash and old_hash in created_by_hash:
            created_item = created_by_hash[old_hash]
            renamed.append({"path": created_item["path"], "old_path": deleted_item["path"], "status": "renamed"})

    large_changed = [item for item in created + modified if item.get("is_large")]
    return {"created": created, "modified": modified, "deleted": deleted, "renamed": renamed, "large_changed": large_changed}


def inspect_workbook(path: Path) -> dict[str, Any]:
    """Safely inspect workbook structure without dumping cell contents."""
    result = {
        "workbook": path.name,
        "path": path.as_posix(),
        "available": False,
        "sheet_names": [],
        "sheet_dimensions": {},
        "expected_sheets_exist": False,
        "missing_expected_sheets": EXPECTED_WORKBOOK_SHEETS[:],
        "eoat_inventory_exists": False,
        "important_inventory_headers_exist": False,
        "missing_important_inventory_headers": IMPORTANT_INVENTORY_COLUMNS[:],
        "data_rows": {},
        "note": "",
    }
    try:
        from openpyxl import load_workbook
    except ImportError:
        result["note"] = "openpyxl not installed; workbook structure was not inspected."
        return result

    try:
        workbook = load_workbook(path, read_only=True, data_only=False)
    except Exception as exc:  # noqa: BLE001 - workbook files can fail for many harmless reasons.
        result["note"] = f"Workbook could not be inspected safely: {exc}"
        return result

    result["available"] = True
    result["sheet_names"] = workbook.sheetnames
    result["missing_expected_sheets"] = [sheet for sheet in EXPECTED_WORKBOOK_SHEETS if sheet not in workbook.sheetnames]
    result["expected_sheets_exist"] = not result["missing_expected_sheets"]
    result["eoat_inventory_exists"] = "EOAT Inventory" in workbook.sheetnames

    for sheet_name in workbook.sheetnames:
        ws = workbook[sheet_name]
        result["sheet_dimensions"][sheet_name] = {"rows": ws.max_row, "columns": ws.max_column}
        result["data_rows"][sheet_name] = count_meaningful_data_rows(ws)

    if result["eoat_inventory_exists"]:
        ws = workbook["EOAT Inventory"]
        headers = [cell.value for cell in next(ws.iter_rows(min_row=1, max_row=1))]
        result["missing_important_inventory_headers"] = [
            header for header in IMPORTANT_INVENTORY_COLUMNS if header not in headers
        ]
        result["important_inventory_headers_exist"] = not result["missing_important_inventory_headers"]
    workbook.close()
    return result


def count_meaningful_data_rows(ws) -> int:
    """Count rows with real entered values, ignoring setup notes and formulas."""
    count = 0
    for row in ws.iter_rows(min_row=2, values_only=True):
        values = [value for value in row if value not in (None, "")]
        if not values:
            continue
        if all(isinstance(value, str) and (value.startswith("Last Updated:") or value.startswith("=")) for value in values):
            continue
        count += 1
    return count


def changed_paths_from_snapshot(snapshot_compare: dict[str, list[dict[str, Any]]]) -> list[dict[str, str]]:
    """Flatten snapshot changes."""
    items = []
    for key, status in (("created", "added"), ("modified", "modified"), ("deleted", "deleted"), ("renamed", "renamed")):
        for item in snapshot_compare.get(key, []):
            items.append({"path": item["path"], "status": status})
    return items


def safe_snippets(project_root: Path, changed_files: list[dict[str, str]]) -> list[dict[str, Any]]:
    """Collect tiny snippets from safe text-like changed files only."""
    snippets = []
    for item in changed_files:
        relative = item["path"]
        path = project_root / relative
        lower_name = path.name.lower()
        if any(part in lower_name for part in SECRET_NAME_PARTS):
            continue
        if path.suffix.lower() not in TEXT_SNIPPET_EXTS or not path.exists():
            continue
        try:
            if path.stat().st_size > SNIPPET_SIZE_LIMIT:
                continue
            lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
        except OSError:
            continue
        snippets.append({"path": relative, "lines": lines[:5]})
    return snippets


def toolkit_suggestions(report_date: str) -> list[str]:
    """Suggest completed items for toolkit files beside the project folder."""
    suggestions = []
    for filename, text in (
        ("setup_eoat_project.py", "Updated the EOAT project setup script"),
        ("daily_status_summary.py", "Updated the daily status summary reporting tool"),
        ("run_daily_status.ps1", "Updated the PowerShell helper script for daily status reporting"),
        ("USAGE.md", "Updated project usage documentation"),
    ):
        path = SCRIPT_DIR / filename
        if path.exists() and datetime.fromtimestamp(path.stat().st_mtime).date().isoformat() == report_date:
            suggestions.append(text)
    return suggestions


def suggested_items_from_activity(project_root: Path, activity: dict[str, Any], report_date: str) -> list[str]:
    """Create plain-language completed suggestions from project activity."""
    paths = [item["path"] for item in activity.get("created_files", []) + activity.get("modified_files", [])]
    created = {item["path"] for item in activity.get("created_files", [])}
    modified = {item["path"] for item in activity.get("modified_files", [])}
    suggestions = []

    if any(path.endswith("EOAT_Master_Tracker.xlsx") for path in created):
        suggestions.append("Created EOAT_Master_Tracker.xlsx with the main audit workbook tabs")
    if any(path.endswith("EOAT_Master_Tracker.xlsx") for path in modified):
        suggestions.append("Updated the EOAT master tracker workbook")
    if "README.md" in created:
        suggestions.append("Created README.md project documentation")
    if "README.md" in modified or "USAGE.md" in modified:
        suggestions.append("Updated project documentation")
    if any("Cell_Photos/" in path and path in created for path in paths):
        suggestions.append("Added EOAT/cell photos to the project folder")
    if any("Raw_Notes/" in path and path in created for path in paths):
        suggestions.append("Added raw EOAT audit notes")
    if any(path.startswith("03_Standards/") for path in paths):
        suggestions.append("Updated draft EOAT standard or PM documentation")
    if any(path.startswith("04_FMEA/") for path in paths):
        suggestions.append("Updated FMEA working files")
    if any(path.startswith("05_Pilot_Project/") for path in paths):
        suggestions.append("Updated pilot project planning files")
    if any(path.startswith("02_KPI_Data/") for path in paths):
        suggestions.append("Updated KPI/baseline data files")
    if any(path.endswith("project_schedule_week1.json") for path in created | modified):
        suggestions.append("Created or updated the Week 1 schedule tracking file")
    if any(path.endswith("task_progress_week1.json") for path in created | modified):
        suggestions.append("Updated Week 1 task progress tracking")
    if project_root.exists() and (project_root / "00_Project_Admin").exists() and (project_root / "01_EOAT_Audit").exists():
        if any(path in created for path in ("README.md", "01_EOAT_Audit/EOAT_Audit_Database/EOAT_Master_Tracker.xlsx")):
            suggestions.append("Created EOAT_Standardization_Project folder structure")

    for check in activity.get("workbook_checks", []):
        if check.get("expected_sheets_exist") and check.get("important_inventory_headers_exist"):
            suggestions.append("Confirmed the EOAT master tracker has the expected tabs and key inventory headers")

    suggestions.extend(toolkit_suggestions(report_date))
    return dedupe(suggestions)


def build_activity_summary(
    project_root: Path,
    report_date: str,
    include_git: bool,
    include_snapshot: bool,
    include_diff_snippets: bool,
) -> dict[str, Any]:
    """Build the internal activity summary object used by the report."""
    summary: dict[str, Any] = {
        "created_files": [],
        "modified_files": [],
        "deleted_files": [],
        "renamed_files": [],
        "commits_today": [],
        "staged_files": [],
        "uncommitted_files": [],
        "suggested_completed_items": [],
        "workbook_checks": [],
        "large_files_changed": [],
        "markdown_report_files_changed": [],
        "scripts_changed": [],
        "diff_snippets": [],
        "notes": [],
        "used_git": False,
        "used_snapshot": False,
    }

    git_info = collect_git_activity(project_root, report_date) if include_git or git_root(project_root) else {"used": False}
    if git_info.get("used"):
        summary["used_git"] = True
        parsed = git_info["parsed_status"]
        summary["created_files"].extend(parsed["created_files"])
        summary["modified_files"].extend(parsed["modified_files"])
        summary["deleted_files"].extend(parsed["deleted_files"])
        summary["renamed_files"].extend(parsed["renamed_files"])
        summary["staged_files"].extend(parsed["staged_files"])
        summary["uncommitted_files"].extend(parsed["uncommitted_files"])
        summary["commits_today"].extend(git_info.get("commits_today", []))
    elif include_git:
        summary["notes"].append(git_info.get("note", "Git activity was requested but no Git repository was available."))

    if include_snapshot or not summary["used_git"]:
        previous = load_json(latest_snapshot_path(project_root), {}) if latest_snapshot_path(project_root).exists() else None
        current = build_snapshot(project_root)
        comparison = compare_snapshots(previous, current)
        summary["used_snapshot"] = True
        summary["created_files"].extend({"path": item["path"], "status": "added"} for item in comparison["created"])
        summary["modified_files"].extend({"path": item["path"], "status": "modified"} for item in comparison["modified"])
        summary["deleted_files"].extend({"path": item["path"], "status": "deleted"} for item in comparison["deleted"])
        summary["renamed_files"].extend({"path": item["path"], "old_path": item.get("old_path", ""), "status": "renamed"} for item in comparison["renamed"])
        summary["large_files_changed"].extend({"path": item["path"], "status": item["status"]} for item in comparison["large_changed"])
        if not previous:
            summary["notes"].append("No previous file snapshot was found; current files were treated as newly detected.")
        summary["_current_snapshot"] = current

    summary["created_files"] = unique_file_items(summary["created_files"])
    summary["modified_files"] = unique_file_items(summary["modified_files"])
    summary["deleted_files"] = unique_file_items(summary["deleted_files"])
    summary["renamed_files"] = unique_file_items(summary["renamed_files"])
    summary["staged_files"] = unique_file_items(summary["staged_files"])
    summary["uncommitted_files"] = unique_file_items(summary["uncommitted_files"])

    changed_files = (
        summary["created_files"]
        + summary["modified_files"]
        + summary["deleted_files"]
        + summary["renamed_files"]
        + summary["staged_files"]
        + summary["uncommitted_files"]
    )
    changed_paths = {item["path"] for item in changed_files}
    workbook_paths = sorted(
        path for path in changed_paths | {"01_EOAT_Audit/EOAT_Audit_Database/EOAT_Master_Tracker.xlsx"} if path.lower().endswith(".xlsx")
    )
    for relative in workbook_paths:
        workbook_path = project_root / relative
        if workbook_path.exists():
            summary["workbook_checks"].append(inspect_workbook(workbook_path))

    summary["markdown_report_files_changed"] = [
        item for item in changed_files if item["path"].lower().endswith(".md")
    ]
    summary["scripts_changed"] = [
        item for item in changed_files if item["path"].lower().endswith((".py", ".ps1"))
    ]
    if include_diff_snippets:
        summary["diff_snippets"] = safe_snippets(project_root, changed_files)

    summary["suggested_completed_items"] = suggested_items_from_activity(project_root, summary, report_date)
    return summary


def unique_file_items(items: Iterable[dict[str, Any]]) -> list[dict[str, Any]]:
    """Dedupe file item dicts by path/status."""
    seen = set()
    result = []
    for item in items:
        key = (item.get("path", ""), item.get("status", ""), item.get("git_status", ""))
        if item.get("path") and key not in seen:
            seen.add(key)
            result.append(dict(item))
    return result


def task_suggestions(activity: dict[str, Any], manual_completed: list[str]) -> dict[str, dict[str, Any]]:
    """Suggest task status updates from activity and manual completed notes."""
    suggestions: dict[str, dict[str, Any]] = {}
    all_evidence = " ".join(activity.get("suggested_completed_items", []) + manual_completed)
    workbook_checks = activity.get("workbook_checks", [])
    workbook_good = any(check.get("expected_sheets_exist") and check.get("important_inventory_headers_exist") for check in workbook_checks)
    issue_rows = max((check.get("data_rows", {}).get("Issue Log", 0) for check in workbook_checks), default=0)
    interview_rows = max((check.get("data_rows", {}).get("Interview Notes", 0) for check in workbook_checks), default=0)
    pilot_rows = max((check.get("data_rows", {}).get("Pilot Candidates", 0) for check in workbook_checks), default=0)
    changed_paths = [item["path"] for key in ("created_files", "modified_files") for item in activity.get(key, [])]

    def suggest(task_text: str, status: str, evidence: str) -> None:
        suggestions[task_text] = {"status": status, "evidence": evidence}

    if "EOAT_Master_Tracker.xlsx" in all_evidence or any(path.endswith("EOAT_Master_Tracker.xlsx") for path in changed_paths):
        suggest("Start EOAT audit database", "Complete", "EOAT master tracker workbook activity detected.")
    if workbook_good:
        suggest("Finalize audit template", "Complete", "Workbook has expected tabs and important EOAT Inventory headers.")
    if any(path == "README.md" or path.endswith("README.md") for path in changed_paths) or "documentation" in all_evidence.lower():
        suggest("Create project folder/document structure", "Complete", "Project documentation/folder setup activity detected.")
        suggest("Review EOAT project scope and deliverables", "Complete", "Project documentation activity detected.")
    if any(path.startswith("01_EOAT_Audit/Cell_Photos/") for path in changed_paths):
        suggest("Decide photo naming system", "In progress", "Cell photo folder activity detected.")
        suggest("Capture overall EOAT photos", "In progress", "Cell photo activity detected.")
        suggest("Capture vacuum cup/gripper photos", "In progress", "Cell photo activity detected.")
        suggest("Capture tubing routing photos", "In progress", "Cell photo activity detected.")
        suggest("Capture sensor and quick disconnect photos", "In progress", "Cell photo activity detected.")
    if issue_rows > 0:
        suggest("Start Issue Log tab", "Complete", "Issue Log contains rows beyond the header.")
    if interview_rows > 0:
        suggest("Interview operator or technician for each audited cell", "In progress", "Interview Notes contains rows beyond the header.")
    if pilot_rows > 0:
        suggest("Flag possible pilot candidate cells", "In progress", "Pilot Candidates contains rows beyond the header.")
    if any("Week1" in path and "Summary" in path for path in changed_paths):
        suggest("Create Week 1 summary", "Complete", "Week 1 summary file activity detected.")

    for text in manual_completed + activity.get("suggested_completed_items", []):
        for task_text in list(suggestions):
            if keyword_match(task_text, text):
                suggestions[task_text]["status"] = "Complete"
                suggestions[task_text]["evidence"] += f" Manual/activity note: {text}"
    return suggestions


def update_task(task: dict[str, Any], status: str, report_date: str, evidence: str = "", notes: str = "") -> None:
    """Update one task progress record."""
    if status not in ALLOWED_STATUSES:
        return
    task["status"] = status
    if status == "Complete" and not task.get("completed_date"):
        task["completed_date"] = report_date
    if status != "Complete":
        task["completed_date"] = ""
    if evidence:
        current = task.get("evidence", [])
        if isinstance(current, str):
            current = [current]
        if evidence not in current:
            current.append(evidence)
        task["evidence"] = current
    if notes:
        task["notes"] = notes


def collect_manual_completed(args: argparse.Namespace, day: int, auto_items: list[str]) -> list[str]:
    """Collect manual completed items and optionally include defaults."""
    manual = normalize_list(args.completed)
    if args.use_day1_defaults and day == 1:
        manual.extend(DAY1_COMPLETED)
    if args.interactive:
        if auto_items:
            print_items("Auto-detected completed items", auto_items)
        use_auto = ask_yes_no("Do you want to use the auto-detected completed items?", True) if auto_items else False
        completed = auto_items[:] if use_auto else []
        completed.extend(manual)
        if day == 1 and ask_yes_no("Include the Day 1 default completed items?", not completed):
            completed.extend(DAY1_COMPLETED)
        if ask_yes_no("Do you want to add additional completed items?", True):
            completed.extend(collect_list("Additional completed"))
        return dedupe(completed)
    return dedupe(auto_items + manual)


def collect_list(section_name: str, defaults: list[str] | None = None) -> list[str]:
    """Collect a bullet list interactively."""
    if defaults:
        print_items(f"Suggested {section_name}", defaults)
        choice = ask("Press Enter to accept, type E to edit, or type N to enter your own", "")
        if choice.lower() == "":
            return defaults
        if choice.lower() == "e":
            edited = [ask(f"{section_name} item {index}", item) for index, item in enumerate(defaults, start=1)]
            return [item for item in edited if item]

    print(f"\nEnter {section_name} items one at a time. Press Enter on a blank line when done.")
    items: list[str] = []
    while True:
        item = input(f"{section_name} item {len(items) + 1}: ").strip()
        if not item:
            break
        items.append(item)
    return items


def choose_status(default: str) -> str:
    """Prompt for one task status."""
    choices = {"c": "Complete", "i": "In progress", "b": "Blocked", "n": "Not started", "s": "Skipped"}
    prompt = "Status: C=Complete, I=In progress, B=Blocked, N=Not started, S=Skipped"
    default_key = next((key for key, value in choices.items() if value == default), "n")
    while True:
        answer = ask(prompt, default_key.upper()).lower()[:1]
        if answer in choices:
            return choices[answer]
        print("Please choose C, I, B, N, or S.")


def update_progress_interactively(
    progress: dict,
    day: int,
    report_date: str,
    suggestions: dict[str, dict[str, Any]],
) -> None:
    """Ask the user to mark current-day tasks and confirm future task suggestions."""
    current_tasks = [task for task in progress.get("tasks", []) if task["day"] == day]
    print_items(f"Planned Day {day} tasks", [task["task_text"] for task in current_tasks])
    for task in current_tasks:
        suggested = suggestions.get(task["task_text"], {})
        default = suggested.get("status", task.get("status", "Not started"))
        print(f"\nTask: {task['task_text']}")
        if suggested:
            print(f"Suggested: {default} ({suggested.get('evidence', 'activity detected')})")
        status = choose_status(default)
        update_task(task, status, report_date, suggested.get("evidence", "Manual daily status confirmation."))

    future_suggestions = [
        (task, suggestions[task["task_text"]])
        for task in progress.get("tasks", [])
        if task["day"] > day and task["task_text"] in suggestions
    ]
    if future_suggestions:
        print("\nFuture scheduled tasks may already have progress:")
    for task, suggestion in future_suggestions:
        print(f"- Day {task['day']}: {task['task_text']} -> suggested {suggestion['status']}")
        if ask_yes_no("Mark this future task with the suggested status?", suggestion["status"] == "Complete"):
            update_task(task, suggestion["status"], report_date, suggestion.get("evidence", "Future task activity detected."))


def update_progress_noninteractive(
    progress: dict,
    day: int,
    report_date: str,
    suggestions: dict[str, dict[str, Any]],
    manual_completed: list[str],
) -> None:
    """Update progress without prompts using conservative suggestions."""
    completed_text = " ".join(manual_completed)
    for task in progress.get("tasks", []):
        suggested = suggestions.get(task["task_text"])
        if task["day"] <= day and keyword_match(task["task_text"], completed_text):
            update_task(task, "Complete", report_date, "Manual completed item matched this scheduled task.")
        elif suggested and task["day"] <= day:
            update_task(task, suggested["status"], report_date, suggested.get("evidence", "Activity matched this scheduled task."))
        elif suggested and task["day"] > day and suggested["status"] == "Complete":
            update_task(task, "Complete", report_date, suggested.get("evidence", "Obvious future task completion detected."))


def task_counts(progress: dict, day: int) -> dict[str, int]:
    """Summarize task progress for the report."""
    current = [task for task in progress.get("tasks", []) if task["day"] == day]
    next_day = [task for task in progress.get("tasks", []) if task["day"] == day + 1]
    carryover = [
        task for task in progress.get("tasks", [])
        if task["day"] <= day and task.get("status") not in ("Complete", "Skipped")
    ]
    blocked = [task for task in progress.get("tasks", []) if task.get("status") == "Blocked"]
    return {
        "current_complete": sum(1 for task in current if task.get("status") == "Complete"),
        "current_total": len(current),
        "next_complete": sum(1 for task in next_day if task.get("status") == "Complete"),
        "next_total": len(next_day),
        "carryover": len(carryover),
        "blocked": len(blocked),
    }


def build_next_day_plan(progress: dict, day: int, manual_plan: list[str]) -> dict[str, list[str]]:
    """Build schedule-aware next-day plan sections."""
    next_day = day + 1
    tasks = progress.get("tasks", [])
    carryover_tasks = [
        task["task_text"]
        for task in tasks
        if task["day"] <= day and task.get("status") not in ("Complete", "Skipped")
    ]
    next_tasks = [
        task["task_text"]
        for task in tasks
        if task["day"] == next_day and task.get("status") not in ("Complete", "Skipped")
    ]
    next_day_tasks = [task for task in tasks if task["day"] == next_day]
    next_complete = sum(1 for task in next_day_tasks if task.get("status") == "Complete")

    later_tasks = [
        task["task_text"]
        for task in tasks
        if task["day"] > next_day and task.get("status") not in ("Complete", "Skipped")
    ]

    plan = {"Carryover": [], "Scheduled Next Steps": [], "Optional Stretch": []}
    plan["Carryover"] = carryover_tasks[:3]
    remaining_slots = max(3, 6 - len(plan["Carryover"]))
    plan["Scheduled Next Steps"] = next_tasks[:remaining_slots]

    current_day_complete = not carryover_tasks
    ahead_on_next_day = next_complete > 0 or (next_day_tasks and next_complete >= max(1, len(next_day_tasks) // 2))
    if current_day_complete and ahead_on_next_day:
        plan["Optional Stretch"] = later_tasks[:2]

    if manual_plan:
        plan["Scheduled Next Steps"].extend(manual_plan)

    used = set()
    for key in plan:
        cleaned = []
        for item in plan[key]:
            item_key = slug(item)
            if item_key not in used:
                used.add(item_key)
                cleaned.append(item)
        plan[key] = cleaned
    return plan


def activity_counts(activity: dict[str, Any]) -> dict[str, Any]:
    """Return concise report counts."""
    return {
        "created": len(activity.get("created_files", [])),
        "modified": len(activity.get("modified_files", [])),
        "deleted": len(activity.get("deleted_files", [])),
        "commits": len(activity.get("commits_today", [])),
        "staged": len(activity.get("staged_files", [])),
        "uncommitted": len(activity.get("uncommitted_files", [])),
    }


def summarize_file_list(activity: dict[str, Any], limit: int = 12) -> list[str]:
    """Create short key changed file lines."""
    items = []
    for key, label in (
        ("created_files", "added"),
        ("modified_files", "modified"),
        ("deleted_files", "deleted"),
        ("renamed_files", "renamed"),
        ("staged_files", "staged"),
        ("uncommitted_files", "uncommitted"),
    ):
        for item in activity.get(key, []):
            path = item.get("path", "")
            if path:
                items.append(f"{path} - {label}")
    return dedupe(items)[:limit]


def bullet_lines(items: list[str]) -> str:
    """Convert list items to Markdown bullets."""
    return "\n".join(f"- {item}" for item in items) if items else "- None"


def plan_markdown(plan: dict[str, list[str]]) -> str:
    """Render next-day plan sections."""
    parts = []
    for heading in ("Carryover", "Scheduled Next Steps", "Optional Stretch"):
        items = plan.get(heading, [])
        if items:
            parts.append(f"{heading}:\n{bullet_lines(items)}")
    return "\n\n".join(parts) if parts else "- No next-day tasks identified"


def build_markdown(
    week: int,
    day: int,
    report_date: str,
    completed: list[str],
    need: list[str],
    next_plan: dict[str, list[str]],
    include_activity_section: bool,
    activity: dict[str, Any],
    notes: list[str],
    progress_counts: dict[str, int],
) -> str:
    """Build the final supervisor-friendly Markdown report."""
    lines = [
        f"Week {week} Day {day} Summary",
        f"Date: {report_date}",
        "",
        "Completed:",
        bullet_lines(completed),
        "",
        "Need:",
        bullet_lines(need),
        "",
        f"Day {day + 1} Plan:",
        plan_markdown(next_plan),
        "",
        "Schedule Progress:",
        f"- Day {day} planned tasks complete: {progress_counts['current_complete']}/{progress_counts['current_total']}",
        f"- Day {day + 1} planned tasks complete early: {progress_counts['next_complete']}/{progress_counts['next_total']}",
        f"- Carryover tasks: {progress_counts['carryover']}",
        f"- Blocked tasks: {progress_counts['blocked']}",
    ]

    if include_activity_section:
        counts = activity_counts(activity)
        staged_summary = f"{counts['staged']} staged file(s)" if counts["staged"] else "None detected"
        uncommitted_summary = f"{counts['uncommitted']} uncommitted file(s)" if counts["uncommitted"] else "None detected"
        lines.extend(
            [
                "",
                "Project Repository Changes:",
                f"- Files created: {counts['created']}",
                f"- Files modified: {counts['modified']}",
                f"- Files deleted: {counts['deleted']}",
                f"- Git commits today: {counts['commits'] if counts['commits'] else 'None detected'}",
                f"- Staged changes: {staged_summary}",
                f"- Uncommitted changes: {uncommitted_summary}",
                "",
                "Key Changed Files:",
                bullet_lines(summarize_file_list(activity)),
            ]
        )

    all_notes = dedupe(notes + activity.get("notes", []))
    if activity.get("diff_snippets"):
        all_notes.append("Diff snippets were included only from safe, small text-like files.")
    if all_notes:
        lines.extend(["", "Notes:", bullet_lines(all_notes)])

    if activity.get("diff_snippets"):
        lines.extend(["", "Safe Text Snippets:"])
        for snippet in activity["diff_snippets"][:5]:
            lines.append(f"- {snippet['path']}")
            for text_line in snippet["lines"]:
                lines.append(f"  {text_line[:160]}")
    return "\n".join(lines) + "\n"


def confirm_overwrite(path: Path, assume_yes: bool = False) -> bool:
    """Return True when it is okay to overwrite an existing report."""
    if not path.exists():
        return True
    if assume_yes:
        return True
    answer = ask(f"{path.name} already exists. Overwrite it? Type YES to confirm", "")
    return answer == "YES"


def create_docx(path: Path, markdown_text: str) -> None:
    """Create a simple DOCX copy when python-docx is available."""
    try:
        from docx import Document
    except ImportError:
        print("python-docx is not installed. Markdown report was still created.")
        return

    document = Document()
    for line in markdown_text.splitlines():
        if not line:
            continue
        if line.startswith("Week ") and " Summary" in line:
            document.add_heading(line, level=1)
        elif line.endswith(":") and not line.startswith("-"):
            document.add_paragraph(line)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            document.add_paragraph(line)
    document.save(path)


def init_snapshot(project_root: Path, report_date: str) -> None:
    """Create a baseline snapshot and exit."""
    snapshot = build_snapshot(project_root)
    path = save_snapshot(project_root, snapshot, report_date)
    print(f"Created baseline snapshot: {path}")


def main() -> None:
    """Create exactly one daily status report."""
    args = parse_args()
    if args.docx:
        args.output_format = "docx"
    project_root = Path(args.project_root).expanduser()
    if not project_root.is_absolute():
        project_root = (SCRIPT_DIR / project_root).resolve()
    else:
        project_root = project_root.resolve()
    report_date = validate_date(args.report_date)

    if args.init_snapshot:
        init_snapshot(project_root, report_date)
        return

    day = args.day if args.day is not None else ask_int("Project day number", 1)
    schedule = ensure_schedule(project_root, args.week)
    progress = ensure_progress(project_root, args.week, schedule)

    activity = build_activity_summary(
        project_root=project_root,
        report_date=report_date,
        include_git=args.include_git,
        include_snapshot=args.include_snapshot,
        include_diff_snippets=args.include_diff_snippets,
    )
    completed = collect_manual_completed(args, day, activity["suggested_completed_items"])
    task_status_suggestions = task_suggestions(activity, completed)

    if args.interactive:
        update_progress_interactively(progress, day, report_date, task_status_suggestions)
    else:
        update_progress_noninteractive(progress, day, report_date, task_status_suggestions, completed)

    if args.use_day1_defaults and day == 1:
        default_need = DAY1_NEED
    else:
        default_need = []
    need = normalize_list(args.need)
    if args.interactive:
        need = collect_list("Need", default_need)
    elif not need:
        need = default_need

    manual_plan = normalize_list(args.plan)
    if args.interactive and ask_yes_no("Do you want to add manual next-day plan items?", False):
        manual_plan.extend(collect_list(f"Day {day + 1} Plan additions"))

    notes = normalize_list(args.note)
    if args.interactive and ask_yes_no("Add optional notes?", False):
        notes.extend(collect_list("Notes"))

    next_plan = build_next_day_plan(progress, day, manual_plan)
    progress_counts = task_counts(progress, day)
    include_activity_section = bool(args.include_git or args.include_snapshot or activity["used_git"] or activity["used_snapshot"])
    if args.interactive:
        include_activity_section = ask_yes_no("Do you want to include a repository activity section in the report?", True)

    output_directory = report_dir(project_root)
    output_directory.mkdir(parents=True, exist_ok=True)
    report_path = output_directory / f"Week{args.week}_Day{day}_Status_{report_date}.md"
    markdown_text = build_markdown(
        args.week,
        day,
        report_date,
        completed,
        need,
        next_plan,
        include_activity_section,
        activity,
        notes,
        progress_counts,
    )

    if not confirm_overwrite(report_path, args.yes):
        print("No report was overwritten. Exiting.")
        return

    report_path.write_text(markdown_text, encoding="utf-8")
    save_json(progress_path(project_root, args.week), progress)
    if activity.get("_current_snapshot"):
        save_snapshot(project_root, activity["_current_snapshot"], report_date)

    if args.output_format == "docx":
        docx_path = report_path.with_suffix(".docx")
        if confirm_overwrite(docx_path, args.yes):
            create_docx(docx_path, markdown_text)

    summary_path = output_directory / f"Week{args.week}_Day{day}_Activity_Summary_{report_date}.json"
    public_summary = {key: value for key, value in activity.items() if not key.startswith("_") and key != "raw"}
    save_json(summary_path, public_summary)
    print(f"Created daily status report: {report_path}")
    print(f"Saved activity summary: {summary_path}")


if __name__ == "__main__":
    main()
